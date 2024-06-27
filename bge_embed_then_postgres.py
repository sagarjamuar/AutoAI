# import psycopg2
# import os
# import fitz  # PyMuPDF
# import docx
# from transformers import AutoTokenizer, AutoModel
# from langchain.text_splitter import CharacterTextSplitter
# import torch
#
# # Load model and tokenizer
# tokenizer = AutoTokenizer.from_pretrained("BAAI/bge-large-en-v1.5")
# model = AutoModel.from_pretrained("BAAI/bge-large-en-v1.5")
# device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
# model.to(device)
#
# # Database connection
# conn = psycopg2.connect(
#     dbname="embeddings_db",
#     user="postgres",
#     password="postgres",
#     host="localhost"
# )
# cur = conn.cursor()
#
# # Function to process DOCX files
# def process_docx(file_path):
#     doc = docx.Document(file_path)
#     return '\n'.join([para.text for para in doc.paragraphs])
#
# # Function to process PDF files
# def process_pdf(file_path):
#     doc = fitz.open(file_path)
#     text = ""
#     for page in doc:
#         text += page.get_text()
#     return text
#
# # Function to get text chunks
# def get_text_chunks(text):
#     text_splitter = CharacterTextSplitter(
#         separator="\n",
#         chunk_size=600,
#         chunk_overlap=50,
#         length_function=len
#     )
#     return text_splitter.split_text(text)
#
# # Function to process files
# def process_files(folder):
#     for root, _, files in os.walk(folder):
#         for file_name in files:
#             if file_name == "Modern_Vehicle_Design.pdf":
#                 print(f"Skipping file: {file_name}")
#                 continue
#             file_path = os.path.join(root, file_name)
#             print(f"Processing file: {file_path}")
#             if file_name.endswith('.pdf'):
#                 text = process_pdf(file_path)
#             elif file_name.endswith('.docx'):
#                 text = process_docx(file_path)
#             else:
#                 continue
#             chunks = get_text_chunks(text)
#             for chunk in chunks:
#                 print(f"Processing chunk of size {len(chunk)}")
#                 inputs = tokenizer(chunk, return_tensors="pt", truncation=True, padding="max_length", max_length=512)
#                 inputs = {k: v.to(device) for k, v in inputs.items()}  # Move inputs to the same device as the model
#                 with torch.no_grad():
#                     outputs = model(**inputs)
#                 embedding = outputs.pooler_output.squeeze().cpu().tolist()
#                 # Insert into PostgreSQL
#                 cur.execute(
#                     "INSERT INTO documents (text, source, embedding) VALUES (%s, %s, %s)",
#                     (chunk, file_path, embedding)
#                 )
#                 conn.commit()
#                 print(f"Inserted embedding of length {len(embedding)} for chunk")
#
# # Directories to process
# for folder in ['/home/ubuntu/Vehicle_Info/INFO', '/home/ubuntu/Vehicle_Info/SPECS']:
#     print(f"Processing folder: {folder}")
#     process_files(folder)
#
# # Close database connection
# cur.close()
# conn.close()

#seperator


import psycopg2
import os
import fitz  # PyMuPDF
import docx
from transformers import AutoTokenizer, AutoModel
from langchain.text_splitter import CharacterTextSplitter
import torch

# Load model and tokenizer
tokenizer = AutoTokenizer.from_pretrained("BAAI/bge-large-en-v1.5")
model = AutoModel.from_pretrained("BAAI/bge-large-en-v1.5")
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model.to(device)

# Database connection
conn = psycopg2.connect(
    dbname="embeddings_db",
    user="postgres",
    password="postgres",
    host="localhost"
)
cur = conn.cursor()

# Check if the table exists and clear it if it does
cur.execute("""
    CREATE TABLE IF NOT EXISTS documents (
        id SERIAL PRIMARY KEY,
        text TEXT,
        source TEXT,
        embedding FLOAT8[]
    );
    DELETE FROM documents;
""")
conn.commit()

# Function to process DOCX files
def process_docx(file_path):
    doc = docx.Document(file_path)
    return '\n'.join([para.text for para in doc.paragraphs])

# Function to process PDF files
def process_pdf(file_path):
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

# Function to get text chunks
def get_text_chunks(text):
    text_splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=300,  # Reduced chunk size
        chunk_overlap=50,
        length_function=len
    )
    return text_splitter.split_text(text)

# Function to process files
def process_files(folder):
    for root, _, files in os.walk(folder):
        for file_name in files:
            if file_name == "Modern_Vehicle_Design.pdf":
                print(f"Skipping file: {file_name}")
                continue
            file_path = os.path.join(root, file_name)
            print(f"Processing file: {file_path}")
            if file_name.endswith('.pdf'):
                text = process_pdf(file_path)
            elif file_name.endswith('.docx'):
                text = process_docx(file_path)
            else:
                continue
            chunks = get_text_chunks(text)
            for chunk in chunks:
                print(f"Processing chunk of size {len(chunk)}")
                inputs = tokenizer(chunk, return_tensors="pt", truncation=True, padding="max_length", max_length=512)
                inputs = {k: v.to(device) for k, v in inputs.items()}  # Move inputs to the same device as the model
                with torch.no_grad():
                    outputs = model(**inputs)
                embedding = outputs.last_hidden_state.mean(dim=1).squeeze().cpu().tolist()  # Use mean pooling
                # Insert into PostgreSQL
                cur.execute(
                    "INSERT INTO documents (text, source, embedding) VALUES (%s, %s, %s)",
                    (chunk, file_path, str(embedding))  # Store embedding as a string
                )
                conn.commit()
                print(f"Inserted embedding of length {len(embedding)} for chunk")

# Directories to process
for folder in ['/home/ubuntu/Vehicle_Info/INFO', '/home/ubuntu/Vehicle_Info/SPECS']:
    print(f"Processing folder: {folder}")
    process_files(folder)

# Close database connection
cur.close()
conn.close()


