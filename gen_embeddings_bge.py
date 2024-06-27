# import os
# import fitz  # PyMuPDF
# import docx
# from transformers import AutoTokenizer, AutoModel
# from langchain.text_splitter import CharacterTextSplitter
# import torch
#
# # Load model and tokenizer
# tokenizer = AutoTokenizer.from_pretrained("BAAI/bge-large-en-v1.5")
# model = AutoModel.from_pretrained("BAAI/bge-large-en-v1.5")
# device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
# model.to(device)
#
# # Function to process DOCX files
# def process_docx(file_path):
#     doc = docx.Document(file_path)
#     return '\n'.join([para.text for para in doc.paragraphs])
#
# # Function to process PDF files
# def process_pdf(file_path):
#     doc = fitz.open(file_path)
#     text = ""
#     for page in doc:
#         text += page.get_text()
#     return text
#
# # Function to get text chunks
# def get_text_chunks(text):
#     text_splitter = CharacterTextSplitter(
#         separator="\n",
#         chunk_size=600,
#         chunk_overlap=50,
#         length_function=len
#     )
#     return text_splitter.split_text(text)
#
# # Function to process files
# def process_files(folder):
#     documents = {}
#     for root, _, files in os.walk(folder):
#         for file_name in files:
#             if file_name == "Modern_Vehicle_Design.pdf":  #skip file
#                 print(f"Skipping file: {file_name}")
#                 continue
#             file_path = os.path.join(root, file_name)
#             print(f"Processing file: {file_path}")
#             if file_name.endswith('.pdf'):
#                 text = process_pdf(file_path)
#             elif file_name.endswith('.docx'):
#                 text = process_docx(file_path)
#             else:
#                 continue
#             chunks = get_text_chunks(text)
#             for chunk in chunks:
#                 print(f"Processing chunk of size {len(chunk)}")
#                 inputs = tokenizer(chunk, return_tensors="pt", truncation=True, padding="max_length", max_length=512)
#                 inputs = {k: v.to(device) for k, v in inputs.items()}  # Move inputs to the same device as the model
#                 with torch.no_grad():
#                     outputs = model(**inputs)
#                 embedding = outputs.pooler_output.squeeze().cpu().tolist()
#                 documents[file_name] = {"text": chunk, "source": file_path, "embedding": embedding}
#                 print(f"Generated embedding of length {len(embedding)} for chunk")
#     return documents
#
# # Directories to process
# documents = {}
# for folder in ['/home/ubuntu/Vehicle_Info/INFO', '/home/ubuntu/Vehicle_Info/SPECS']:
#     print(f"Processing folder: {folder}")
#     documents.update(process_files(folder))
#
# # Example output
# for result in list(documents.values())[:3]:  # Display first 3 results for brevity
#     print("Text:", result["text"])
#     print("Source:", result["source"])
#     print("Embedding Length:", len(result["embedding"]))
#     print()


#seperator, below is a working code !

import os
import fitz  # PyMuPDF
import docx
from transformers import AutoTokenizer, AutoModel
from langchain.text_splitter import CharacterTextSplitter
import torch
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity

print("hello")
# Load model and tokenizer
tokenizer = AutoTokenizer.from_pretrained("BAAI/bge-large-en-v1.5")
model = AutoModel.from_pretrained("BAAI/bge-large-en-v1.5")
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model.to(device)

# Function to process DOCX files
def process_docx(file_path):
    doc = docx.Document(file_path)
    return '\n'.join([para.text for para in doc.paragraphs])

# Function to process PDF files
def process_pdf(file_path):
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

# Function to get text chunks
def get_text_chunks(text):
    text_splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=600,
        chunk_overlap=50,
        length_function=len
    )
    return text_splitter.split_text(text)

# Function to process files
def process_files(folder):
    documents = {}
    for root, _, files in os.walk(folder):
        for file_name in files:
            if file_name == "Modern_Vehicle_Design.pdf":  # Skip file
                print(f"Skipping file: {file_name}")
                continue
            file_path = os.path.join(root, file_name)
            print(f"Processing file: {file_path}")
            if file_name.endswith('.pdf'):
                text = process_pdf(file_path)
            elif file_name.endswith('.docx'):
                text = process_docx(file_path)
            else:
                continue
            chunks = get_text_chunks(text)
            for chunk in chunks:
                print(f"Processing chunk of size {len(chunk)}")
                inputs = tokenizer(chunk, return_tensors="pt", truncation=True, padding="max_length", max_length=512)
                inputs = {k: v.to(device) for k, v in inputs.items()}  # Move inputs to the same device as the model
                with torch.no_grad():
                    outputs = model(**inputs)
                embedding = outputs.last_hidden_state.mean(dim=1).squeeze().cpu().tolist()
                if file_name not in documents:
                    documents[file_name] = []
                documents[file_name].append({"text": chunk, "source": file_path, "embedding": embedding})
                print(f"Generated embedding of length {len(embedding)} for chunk")
    return documents

# Directories to process
documents = {}
for folder in ['/home/ubuntu/Vehicle_Info/INFO', '/home/ubuntu/Vehicle_Info/SPECS']:
    print(f"Processing folder: {folder}")
    documents.update(process_files(folder))

# Function to get embedding for a query
def get_embedding(text):
    inputs = tokenizer(text, return_tensors="pt", truncation=True, padding="max_length", max_length=512)
    inputs = {k: v.to(device) for k, v in inputs.items()}
    with torch.no_grad():
        outputs = model(**inputs)
    embedding = outputs.last_hidden_state.mean(dim=1).squeeze().cpu().numpy()
    return embedding

# Query and embed it
# query = "How many seats are there in Renault Kiger?"
query = "bmw 6 series start of production?"
query_embedding = get_embedding(query)

# Calculate cosine similarity
all_embeddings = []
all_sources = []
for file_name, chunks in documents.items():
    for chunk in chunks:
        all_embeddings.append(chunk['embedding'])
        all_sources.append(chunk['source'])

all_embeddings = np.array(all_embeddings)
similarity_scores = cosine_similarity([query_embedding], all_embeddings)[0]

# Get top 3 results
top_indices = similarity_scores.argsort()[-3:][::-1]
top_sources = [all_sources[i] for i in top_indices]

# Print top 3 sources
print("Top 3 sources for the query:")
for i, source in enumerate(top_sources):
    print(f"{i+1}. {source}")


# import os
# import fitz  # PyMuPDF
# import docx
# from transformers import AutoTokenizer, AutoModel
# from langchain.text_splitter import CharacterTextSplitter
# import torch
# import numpy as np
# from sklearn.metrics.pairwise import cosine_similarity
# import psycopg2
#
# print("hello")
# # Load model and tokenizer
# tokenizer = AutoTokenizer.from_pretrained("BAAI/bge-large-en-v1.5")
# model = AutoModel.from_pretrained("BAAI/bge-large-en-v1.5")
# device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
# model.to(device)
#
# # Database connection
# conn = psycopg2.connect(
#     dbname="embeddings_db",
#     user="postgres",
#     password="postgres",
#     host="localhost"
# )
# cur = conn.cursor()
#
# # Check if the dummy table exists and clear it if it does
# cur.execute("""
#     CREATE TABLE IF NOT EXISTS dummy_documents (
#         id SERIAL PRIMARY KEY,
#         source TEXT,
#         embedding TEXT
#     );
#     DELETE FROM dummy_documents;
# """)
# conn.commit()
#
# # Function to process DOCX files
# def process_docx(file_path):
#     doc = docx.Document(file_path)
#     return '\n'.join([para.text for para in doc.paragraphs])
#
# # Function to process PDF files
# def process_pdf(file_path):
#     doc = fitz.open(file_path)
#     text = ""
#     for page in doc:
#         text += page.get_text()
#     return text
#
# # Function to get text chunks
# def get_text_chunks(text):
#     text_splitter = CharacterTextSplitter(
#         separator="\n",
#         chunk_size=600,
#         chunk_overlap=50,
#         length_function=len
#     )
#     return text_splitter.split_text(text)
#
# # Function to process files
# def process_files(folder):
#     documents = {}
#     for root, _, files in os.walk(folder):
#         for file_name in files:
#             if file_name == "Modern_Vehicle_Design.pdf":  # Skip file
#                 print(f"Skipping file: {file_name}")
#                 continue
#             file_path = os.path.join(root, file_name)
#             print(f"Processing file: {file_path}")
#             if file_name.endswith('.pdf'):
#                 text = process_pdf(file_path)
#             elif file_name.endswith('.docx'):
#                 text = process_docx(file_path)
#             else:
#                 continue
#             chunks = get_text_chunks(text)
#             for chunk in chunks:
#                 print(f"Processing chunk of size {len(chunk)}")
#                 inputs = tokenizer(chunk, return_tensors="pt", truncation=True, padding="max_length", max_length=512)
#                 inputs = {k: v.to(device) for k, v in inputs.items()}  # Move inputs to the same device as the model
#                 with torch.no_grad():
#                     outputs = model(**inputs)
#                 embedding = outputs.last_hidden_state.mean(dim=1).squeeze().cpu().tolist()
#                 cur.execute(
#                     "INSERT INTO dummy_documents (source, embedding) VALUES (%s, %s)",
#                     (file_path, str(embedding))
#                 )
#                 conn.commit()
#                 print(f"Generated embedding of length {len(embedding)} for chunk")
#     return documents
#
# # Directories to process
# for folder in ['/home/ubuntu/Vehicle_Info/INFO', '/home/ubuntu/Vehicle_Info/SPECS']:
#     print(f"Processing folder: {folder}")
#     process_files(folder)
#
# # Function to get embedding for a query
# def get_embedding(text):
#     inputs = tokenizer(text, return_tensors="pt", truncation=True, padding="max_length", max_length=512)
#     inputs = {k: v.to(device) for k, v in inputs.items()}
#     with torch.no_grad():
#         outputs = model(**inputs)
#     embedding = outputs.last_hidden_state.mean(dim=1).squeeze().cpu().numpy()
#     return embedding
#
# # Query and embed it
# # query = "How many seats are there in Renault Kiger?"
# query = "How to repair bmw?"
# query_embedding = get_embedding(query)
#
# # Fetch all embeddings from PostgreSQL
# cur.execute("SELECT source, embedding FROM dummy_documents")
# rows = cur.fetchall()
#
# # Calculate cosine similarity
# all_embeddings = []
# all_sources = []
# for row in rows:
#     all_sources.append(row[0])
#     all_embeddings.append(eval(row[1]))  # Convert string back to list
#
# all_embeddings = np.array(all_embeddings)
# similarity_scores = cosine_similarity([query_embedding], all_embeddings)[0]
#
# # Get top 3 results
# top_indices = similarity_scores.argsort()[-3:][::-1]
# top_sources = [all_sources[i] for i in top_indices]
#
# # Print top 3 sources
# print("Top 3 sources for the query:")
# for i, source in enumerate(top_sources):
#     print(f"{i+1}. {source}")
#
# # Drop the dummy_documents table
# cur.execute("DROP TABLE IF EXISTS dummy_documents")
# conn.commit()
#
# # Close database connection
# cur.close()
# conn.close()
